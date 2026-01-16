from pydoc import doc
from docx import Document
from config import get_data_file_path

def find_document(filename: str):
    file_path = get_data_file_path(filename)
    doc = Document(file_path)
    return doc

doc = find_document('essay_FINAL_general.docx')

#1a
def count_paragraphs(doc):
    return len(doc.paragraphs)

#1b
print("Number of paragraphs in the document:", count_paragraphs(doc))
font = doc.styles['Normal'].font

#1c
def extract_file(doc):
    header_lines = []
    title = ""
    body_paragraphs = []

    i = 0
    while i < len(doc.paragraphs) and doc.paragraphs[i].text.strip() != "":
        header_lines.append(doc.paragraphs[i].text.strip())
        i += 1

    while i < len(doc.paragraphs) and doc.paragraphs[i].text.strip() == "":
        i += 1

    # 1e - centered and bold title
    while i < len(doc.paragraphs) and doc.paragraphs[i].text.strip() != "":
        paragraph = doc.paragraphs[i]
        title = paragraph.text
        isBold = paragraph.style.font.bold
        isCentered = paragraph.alignment == 1  # 1 corresponds to centered alignment
        print(f"Title: '{title}', Bold: {isBold}, Centered: {isCentered}")

        i += 1

    for p in doc.paragraphs[i:]:
        t = p.text.strip()
        if t != "":
            body_paragraphs.append(t)

    return "\n".join(header_lines), title, "\n\n".join(body_paragraphs)

extracted = extract_file(doc)
print(f'Header: {extracted[0]}, Title: {extracted[1]}, Body: {extracted[2][:100]}...')